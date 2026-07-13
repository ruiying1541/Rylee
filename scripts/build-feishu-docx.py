from pathlib import Path
import json
import re
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html


WORKSPACE = Path(__file__).resolve().parents[1]
ARTICLES_PATH = WORKSPACE / "data" / "articles.json"
OUT_DIR = WORKSPACE / "deliverables"
SLUG = sys.argv[1] if len(sys.argv) > 1 else "ld5e2kwa3yld1s65"


def main():
    data = json.loads(ARTICLES_PATH.read_text(encoding="utf-8"))
    article = next((item for item in data["articles"] if item["slug"] == SLUG), None)
    if not article:
        raise SystemExit(f"Article not found: {SLUG}")

    OUT_DIR.mkdir(exist_ok=True)
    safe_title = sanitize_filename(article["title"])
    docx_path = OUT_DIR / f"{safe_title}.polished.feishu.docx"
    html_path = OUT_DIR / f"{safe_title}.html"

    doc = Document()
    configure_document(doc)

    add_title(doc, article["title"])
    add_metadata(doc, article)
    add_html_content(doc, article["html"])

    doc.core_properties.title = article["title"]
    doc.core_properties.subject = "Imported from Yuque for Feishu document import"
    doc.core_properties.keywords = "Yuque, Feishu, UI/UX, Claude Skills"
    doc.save(docx_path)

    html_path.write_text(build_import_html(article), encoding="utf-8")
    print(docx_path)
    print(html_path)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.86)
    section.right_margin = Inches(0.92)
    section.bottom_margin = Inches(0.86)
    section.left_margin = Inches(0.92)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Aptos", "PingFang SC", 11.2, "1F2329")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.32

    for name, size, color, before, after in [
        ("Heading 1", 22, "1F2329", 22, 8),
        ("Heading 2", 17, "1F2329", 20, 7),
        ("Heading 3", 15.2, "1F2329", 18, 7),
    ]:
        style = doc.styles[name]
        set_style_font(style, "Aptos Display", "PingFang SC", size, color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2
        style.paragraph_format.keep_with_next = True

    link_style = doc.styles.add_style("Feishu Link", WD_STYLE_TYPE.CHARACTER)
    set_style_font(link_style, "Aptos", "PingFang SC", 10.2, "245BDB")
    link_style.font.color.rgb = RGBColor(17, 85, 204)
    link_style.font.underline = True

    muted_style = doc.styles.add_style("Muted Metadata", WD_STYLE_TYPE.CHARACTER)
    set_style_font(muted_style, "Aptos", "PingFang SC", 9.4, "646A73")

    focus = doc.styles.add_style("Focus Line", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(focus, "Aptos", "PingFang SC", 10.8, "1F2329")
    focus.paragraph_format.left_indent = Inches(0.12)
    focus.paragraph_format.right_indent = Inches(0.08)
    focus.paragraph_format.space_before = Pt(2)
    focus.paragraph_format.space_after = Pt(9)
    focus.paragraph_format.line_spacing = 1.25

    link_para = doc.styles.add_style("Link Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(link_para, "Aptos", "PingFang SC", 10.2, "245BDB")
    link_para.paragraph_format.left_indent = Inches(0.12)
    link_para.paragraph_format.space_after = Pt(3)
    link_para.paragraph_format.line_spacing = 1.15

    intro = doc.styles.add_style("Intro Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(intro, "Aptos", "PingFang SC", 12.0, "1F2329")
    intro.paragraph_format.space_after = Pt(12)
    intro.paragraph_format.line_spacing = 1.32

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    run = footer.add_run("飞书导入精排版")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(143, 149, 158)


def add_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, "Aptos Display", "PingFang SC", 24, "1F2329")
    run.bold = True


def add_metadata(doc, article):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run("来源：").style = "Muted Metadata"
    add_hyperlink(paragraph, article["sourceUrl"], article["sourceUrl"])
    if article.get("updatedAt"):
        paragraph.add_run().add_break(WD_BREAK.LINE)
        paragraph.add_run(f"语雀更新时间：{article['updatedAt']}").style = "Muted Metadata"


def add_html_content(doc, html_text):
    root = html.fragment_fromstring(f"<article>{html_text}</article>", create_parent=True)
    article = root.find(".//article")

    for element in article:
        tag = clean_tag(element.tag)
        if tag in {"h1", "h2", "h3"}:
            level = {"h1": 1, "h2": 2, "h3": 2}[tag]
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            append_inline(paragraph, element)
        elif tag == "p":
            text = normalized_text(element)
            if text:
                style = choose_paragraph_style(element)
                paragraph = doc.add_paragraph(style=style)
                if style == "Focus Line":
                    shade_paragraph(paragraph, "F7F8FA")
                append_inline(paragraph, element)
        elif tag in {"ul", "ol"}:
            add_list(doc, element, ordered=(tag == "ol"))
        elif tag == "blockquote":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            append_inline(paragraph, element)
        else:
            text = normalized_text(element)
            if text:
                doc.add_paragraph(text)


def add_list(doc, list_element, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for li in list_element.xpath("./li"):
        paragraph = doc.add_paragraph(style=style)
        append_inline(paragraph, li)


def append_inline(paragraph, element, inherited_bold=False):
    if element.text:
        add_run(paragraph, element.text, inherited_bold)

    for child in element:
        tag = clean_tag(child.tag)
        bold = inherited_bold or tag in {"strong", "b"}
        if tag == "a":
            text = normalized_text(child) or child.get("href", "")
            href = child.get("href")
            if href:
                add_hyperlink(paragraph, text, href)
            else:
                add_run(paragraph, text, bold)
        elif tag == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        else:
            append_inline(paragraph, child, bold)

        if child.tail:
            add_run(paragraph, child.tail, inherited_bold)


def add_run(paragraph, text, bold=False):
    if not text:
        return
    run = paragraph.add_run(text)
    if bold:
        run.bold = True


def choose_paragraph_style(element):
    text = normalized_text(element)
    if text.startswith("专注领域"):
        return "Focus Line"
    if text.startswith("http://") or text.startswith("https://"):
        return "Link Paragraph"
    return "Normal"


def set_style_font(style, latin, east_asia, size, color):
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)


def set_run_font(run, latin, east_asia, size, color):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FeishuLink")
    properties.append(style)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def normalized_text(element):
    return re.sub(r"\s+", " ", element.text_content()).strip()


def clean_tag(tag):
    return str(tag).lower() if isinstance(tag, str) else ""


def sanitize_filename(value):
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip()[:120]


def build_import_html(article):
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{escape_html(article["title"])}</title>
  <style>
    body {{ font-family: Arial, "PingFang SC", "Microsoft YaHei", sans-serif; line-height: 1.6; max-width: 820px; margin: 40px auto; color: #111; }}
    h1 {{ font-size: 30px; font-weight: 400; margin: 0 0 8px; }}
    h3 {{ font-size: 18px; font-weight: 400; margin: 28px 0 8px; color: #434343; }}
    p {{ margin: 0 0 12px; }}
    a {{ color: #1155cc; }}
    .meta {{ color: #555; font-size: 13px; margin-bottom: 24px; }}
  </style>
</head>
<body>
  <h1>{escape_html(article["title"])}</h1>
  <p class="meta">来源：<a href="{escape_html(article["sourceUrl"])}">{escape_html(article["sourceUrl"])}</a><br>语雀更新时间：{escape_html(article.get("updatedAt", ""))}</p>
  {article["html"]}
</body>
</html>
"""


def escape_html(value):
    return (
        str(value)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


if __name__ == "__main__":
    main()
